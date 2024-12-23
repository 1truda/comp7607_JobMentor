from docx import Document

# Create a Word Document
doc = Document()

# Adding a title
doc.add_heading('AI Agent Proposal for Assisting Students with Job Search', level=1)

# Adding a separator
doc.add_paragraph('---')

# Adding Objective section
doc.add_heading('Objective', level=2)
doc.add_paragraph(
    "The goal of this AI agent is to assist students in job search preparation and streamline the associated processes, "
    "including quality tests, interview practice, scheduling, and coding interview preparation. The system aims to enhance "
    "efficiency and learning outcomes by utilizing AI-driven techniques in areas such as question-answering, interview simulation, "
    "automated scheduling, and coding problem-solving."
)

# Adding a separator
doc.add_paragraph('---')

# Adding Data Sets section
doc.add_heading('Data Sets', level=2)

# Adding individual data sets
data_sets = [
    ("Quality Test Questions", 
     "Sources: Quality test platforms such as NowCoder and GKZenti.\n"
     "Content: Various types of questions, including math reasoning, context reasoning, and office skills assessments."),
    
    ("Email Scheduling Data", 
     "Collected From: Students’ email accounts for personalized scheduling and management.\n"
     "Purpose: To learn patterns in communication and time management."),
    
    ("Quora Question Pairs", 
     "Usage: Model semantic similarity in interview preparation scenarios. Data will be cleaned and augmented to emphasize job-related topics."),
    
    ("Coding Interview Questions", 
     "Scope: Over 3,300 coding problems and solutions, covering data structures, algorithms, and various programming concepts in Python."),
    
    ("Customer-Service Conversations", 
     "Source: JD.com dataset containing 10,000 anonymized sessions, simulating interview dialogues.\n"
     "Additional Data: Frequently asked interview questions and responses for prompt engineering and interview simulation.")
]

for title, content in data_sets:
    doc.add_heading(title, level=3)
    doc.add_paragraph(content)

# Adding a separator
doc.add_paragraph('---')

# Adding Methodology section
doc.add_heading('Methodology', level=2)

# Sub-sections in Methodology
methodology = [
    ("Quality Test Assistance", 
     "QA Agent: Trains on datasets to answer various quality test questions correctly. Uses OCR to extract text from screenshots "
     "and automated browser interaction to complete online tests."),
    
    ("Scheduling System", 
     "Email Parsing & To-Do Lists: Utilizes LLMs to extract and summarize email content, recognize time-related information, and "
     "create optimized schedules. Automates confirmations, reminders, and generates to-do lists.\n"
     "Techniques: Few-shot learning and parameterized prompts for accurate information extraction, combined with a MySQL database "
     "for managing scheduling data."),
    
    ("Interview Preparation", 
     "Semantic Search & Recommendation: Uses BM25, BERT embeddings, and fine-tuned SBERT for question similarity. DPR enhances "
     "retrieval for complex queries. User feedback loops improve the recommendation system.\n"
     "RAG System: Combines information retrieval with generative models to simulate interviews. Retrieves relevant content and "
     "provides tailored feedback. LangChain integrates the modules for dynamic Q&A."),
    
    ("Coding Interview Agent", 
     "QA Agent: Powered by Llama-3.1 and OS-Copilot, generates accurate Python solutions to coding problems. Provides interactive "
     "assistance, teaching techniques, and enhances coding skills."),
    
    ("Simulated Interview System", 
     "Interactive Simulation: Implements Retrieval-Augmented Generation (RAG) for realistic interview scenarios. Uses LangChain for "
     "integration, feedback generation, and managing conversation flow.")
]

for title, content in methodology:
    doc.add_heading(title, level=3)
    doc.add_paragraph(content)

# Adding a separator
doc.add_paragraph('---')

# Adding Technical Framework section
doc.add_heading('Technical Framework', level=2)

# Sub-sections in Technical Framework
frameworks = [
    ("AI Models & Frameworks", 
     "Llama-3.1: Core NLP model for natural language understanding and Python code generation.\n"
     "OS-Copilot: Enhances coding assistance and automates solution suggestions.\n"
     "BM25 & BERT: Baseline and semantic similarity search for interview preparation.\n"
     "RAG & LangChain: Combine retrieval and generative models for an interactive interview experience."),
    
    ("APIs & Databases", 
     "Outlook API: Automates email management and scheduling tasks.\n"
     "MySQL: Securely stores and manages user scheduling data.\n"
     "APIs for Browser Automation: Executes tasks like test completion."),
    
    ("Integration & Automation", 
     "LangChain: Manages context, tracks user interactions, and integrates tools and models for a cohesive interview and learning "
     "experience.\n"
     "Scripts & Automation: Facilitates interaction with external systems (e.g., email, scheduling platforms, test websites).")
]

for title, content in frameworks:
    doc.add_heading(title, level=3)
    doc.add_paragraph(content)

# Adding a separator
doc.add_paragraph('---')

# Adding References section
doc.add_heading('References', level=2)
references = [
    "Retrieval-Augmented Generation for Large Language Models: A Survey Gao et al., 2023.",
    "Interleaving Retrieval with Chain-of-Thought Reasoning for Knowledge-Intensive Multi-Step Questions, ACL 2023",
    "Retrieval Augmented Thoughts Elicit Context-Aware Reasoning in Large Language Models",
    "OS-COPILOT: TOWARDS GENERALIST COMPUTER AGENTS WITH SELF-IMPROVEMENT",
    "OS-pilot",
    "ChatTTS"
]

for ref in references:
    doc.add_paragraph(ref)

# Save the document
file_path = '/mnt/data/AI_Agent_Proposal.docx'
doc.save(file_path)

file_path
